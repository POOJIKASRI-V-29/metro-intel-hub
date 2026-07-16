"""
DOCX Extraction module for the KMRL Document Intelligence Platform.

This module utilizes python-docx to process and extract unstructured paragraph 
text and structured table data from Microsoft Word documents (.docx) via file paths 
or raw binary memory streams.
"""

import io
import logging
from typing import BinaryIO, List
import docx  # type: ignore
from pydantic import BaseModel, Field

# Setup logger mapping to Stage 0 configurations
logger = logging.getLogger("document_intelligence.ingestion.docx_parser")


class ParsedElement(BaseModel):
    """
    Data model representing an extracted structural element from a Word document.
    """
    element_type: str = Field(..., description="The type of element: 'paragraph' or 'table_cell'.")
    text: str = Field(..., description="The raw string content extracted from the element.")
    index: int = Field(..., description="The sequential position index of the element in the document.")


class DocxParser:
    """
    Handles robust extraction of text and inline tabular data from DOCX files.
    """

    def __init__(self) -> None:
        """Initializes the Word document parser component."""
        pass

    def parse_stream(self, file_stream: BinaryIO) -> List[ParsedElement]:
        """
        Parses an in-memory binary stream containing a DOCX file.

        Args:
            file_stream: A seekable file-like binary object containing a Word document.

        Returns:
            A linear list of ParsedElement objects maintaining reading sequence.

        Raises:
            ValueError: If the file stream is invalid or corrupted.
        """
        try:
            file_stream.seek(0)
            stream_bytes = file_stream.read()
            memory_buffer = io.BytesIO(stream_bytes)
            
            # Load the document structure directly from memory buffer
            doc = docx.Document(memory_buffer)
            return self._extract_content(doc)

        except Exception as error:
            logger.exception("Failed to parse DOCX structural content from the provided stream.")
            raise ValueError("The provided stream is not a valid or readable DOCX document.") from error

    def parse_file(self, file_path: str) -> List[ParsedElement]:
        """
        Parses a DOCX file loaded directly from the local file system.

        Args:
            file_path: The absolute or relative path to the target DOCX file.

        Returns:
            A linear list of ParsedElement objects maintaining reading sequence.

        Raises:
            ValueError: If the file path is invalid, unreadable, or corrupted.
        """
        try:
            doc = docx.Document(file_path)
            return self._extract_content(doc)

        except FileNotFoundError:
            logger.error(f"DOCX file not found at path: {file_path}")
            raise ValueError(f"File not found at target path: {file_path}")
        except Exception as error:
            logger.exception(f"Structural layout parsing failure for file at: {file_path}")
            raise ValueError(f"Corrupt or invalid DOCX document layout at: {file_path}") from error

    def _extract_content(self, doc: docx.document.Document) -> List[ParsedElement]:
        """
        Internal worker that extracts text from paragraphs and tables sequentially.

        Args:
            doc: An initialized python-docx Document object.

        Returns:
            A list of ParsedElement records.
        """
        elements: List[ParsedElement] = []
        element_counter = 0

        # Extract standard text paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # Filter out completely empty spacer paragraphs
                elements.append(
                    ParsedElement(
                        element_type="paragraph",
                        text=text,
                        index=element_counter
                    )
                )
                element_counter += 1

        # Extract structured text inside tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                row_texts: List[str] = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                
                # Combine row elements with a clear delimiter for subsequent chunkers
                if row_texts:
                    combined_row_text = " | ".join(row_texts)
                    elements.append(
                        ParsedElement(
                            element_type="table_row",
                            text=f"Table {table_idx + 1}, Row {row_idx + 1}: {combined_row_text}",
                            index=element_counter
                        )
                    )
                    element_counter += 1

        logger.info(f"Successfully extracted {len(elements)} structural elements from DOCX document.")
        return elements